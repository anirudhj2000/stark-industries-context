#!/usr/bin/env python3
"""
DOCX Extraction Helper Script

Extracts text, paragraphs, headings, and tables from DOCX files using python-docx.
"""
import argparse
import json
import sys
from pathlib import Path
from typing import Dict, List, Any

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def extract_paragraph_info(paragraph: Paragraph) -> Dict[str, Any]:
    """Extract information from a paragraph including style and runs"""
    # Determine if it's a heading
    style_name = paragraph.style.name if paragraph.style else "Normal"
    is_heading = style_name.startswith("Heading")
    heading_level = 0

    if is_heading:
        try:
            heading_level = int(style_name.replace("Heading", "").strip())
        except:
            heading_level = 1

    return {
        "text": paragraph.text,
        "style": style_name,
        "is_heading": is_heading,
        "heading_level": heading_level if is_heading else 0,
        "runs": len(paragraph.runs)
    }


def extract_table_info(table: Table) -> Dict[str, Any]:
    """Extract table data as rows and cells"""
    rows_data = []

    for row in table.rows:
        cells_data = []
        for cell in row.cells:
            cells_data.append(cell.text.strip())
        rows_data.append(cells_data)

    # Try to detect if first row is header
    has_header = False
    if len(rows_data) > 1:
        first_row = rows_data[0]
        # Header detection: check if first row cells are non-empty
        has_header = any(cell.strip() for cell in first_row)

    return {
        "rows": len(table.rows),
        "columns": len(table.columns) if table.rows else 0,
        "has_header": has_header,
        "data": rows_data
    }


def extract_docx(file_path: Path) -> Dict[str, Any]:
    """Extract content from DOCX file including paragraphs, headings, and tables"""
    try:
        doc = Document(file_path)

        # Extract paragraphs
        paragraphs = []
        headings = []
        total_words = 0

        for para in doc.paragraphs:
            para_info = extract_paragraph_info(para)
            paragraphs.append(para_info)

            # Count words
            if para_info["text"].strip():
                total_words += len(para_info["text"].split())

            # Track headings separately
            if para_info["is_heading"]:
                headings.append({
                    "level": para_info["heading_level"],
                    "text": para_info["text"]
                })

        # Extract tables
        tables = []
        for table in doc.tables:
            table_info = extract_table_info(table)
            tables.append(table_info)

        # Extract core properties (metadata)
        core_props = doc.core_properties

        result = {
            "file_name": file_path.name,
            "file_size_bytes": file_path.stat().st_size,
            "format": "DOCX",
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs,
            "heading_count": len(headings),
            "headings": headings,
            "table_count": len(tables),
            "tables": tables,
            "total_words": total_words,
            "metadata": {
                "engine": "python-docx",
                "title": core_props.title if core_props.title else "",
                "author": core_props.author if core_props.author else "",
                "subject": core_props.subject if core_props.subject else "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by if core_props.last_modified_by else "",
                "revision": core_props.revision if core_props.revision else 0
            }
        }

        return result

    except Exception as e:
        raise Exception(f"Failed to extract DOCX file: {e}")


def main():
    parser = argparse.ArgumentParser(description='Extract content from DOCX files')
    parser.add_argument('--input', required=True, help='Path to DOCX file')
    parser.add_argument('--output', required=True, help='Output JSON file')

    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: DOCX file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    try:
        result = extract_docx(input_path)

        # Write output
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

        print(f"SUCCESS: Extracted {result['paragraph_count']} paragraphs from {input_path.name}")
        print(f"Total words: {result['total_words']}")
        print(f"Headings: {result['heading_count']}, Tables: {result['table_count']}")

    except Exception as e:
        print(f"ERROR: Extraction failed: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
